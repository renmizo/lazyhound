"""Default DOCX template generation for report builder.

Creates a minimal but professional-looking Word document template that the
report builder uses when no custom template is provided.
"""

from __future__ import annotations

from pathlib import Path


def create_default_template(dest: Path) -> Path:
    """Generate a default DOCX template at *dest*.

    The template contains:
    - A title page placeholder (Heading style)
    - Pre-defined heading styles (Heading 1-3)
    - Normal body text style
    - A table style for findings tables

    Uses python-docx to build the document.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -- Customise built-in styles ----------------------------------------
    style = doc.styles["Title"]
    font = style.font
    font.size = Pt(28)
    font.bold = True
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(4)

    for level, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)]:
        s = doc.styles[level]
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(4)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Calibri"
    normal.paragraph_format.space_after = Pt(6)

    # -- Cover page -------------------------------------------------------
    doc.add_paragraph("LazyHound", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Security Assessment Report")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("{{domain}}  |  {{date}}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # -- Table of Contents placeholder ------------------------------------
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("(auto-generated on report creation)")
    doc.add_page_break()

    # -- Section stubs (removable by builder) -----------------------------
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{section:summary}}")

    doc.add_heading("Findings", level=1)
    doc.add_paragraph("{{section:findings}}")

    doc.add_heading("Attack Paths", level=1)
    doc.add_paragraph("{{section:attack_paths}}")

    doc.add_heading("Delegation Map", level=1)
    doc.add_paragraph("{{section:delegation}}")

    doc.add_heading("Domain Trust Map", level=1)
    doc.add_paragraph("{{section:domain_trust}}")

    # -- Save -------------------------------------------------------------
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


def ensure_templates_dir(base_dir: str | Path) -> Path:
    """Create the templates/ directory under *base_dir* and seed default.docx.

    Returns the path to the templates directory.
    """
    templates_dir = Path(base_dir) / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)

    default_docx = templates_dir / "default.docx"
    if not default_docx.exists():
        create_default_template(default_docx)

    return templates_dir
